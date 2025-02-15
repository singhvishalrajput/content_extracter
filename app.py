from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import tempfile
import PyPDF2
from docx import Document
import pandas as pd
import pytesseract
from PIL import Image
import openpyxl
import csv
import chardet

app = Flask(__name__)
CORS(app)

# Set Tesseract Path - MODIFY THIS PATH according to your installation
# Common Windows paths:
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe'
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Modify this line!

# Configure allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'xlsx', 'csv', 'jpeg', 'jpg', 'png', 'txt'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_from_pdf(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    return text.strip()

def extract_from_docx(file_path):
    text = ""
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text.strip()

def extract_from_excel(file_path):
    df = pd.read_excel(file_path)
    return df.to_string(index=False)

def extract_from_csv(file_path):
    # Detect file encoding
    with open(file_path, 'rb') as file:
        raw_data = file.read()
        result = chardet.detect(raw_data)
        encoding = result['encoding']

    text = ""
    with open(file_path, 'r', encoding=encoding) as file:
        csv_reader = csv.reader(file)
        for row in csv_reader:
            text += ",".join(row) + "\n"
    return text.strip()

def extract_from_image(file_path):
    try:
        # Verify Tesseract installation
        if not os.path.exists(pytesseract.pytesseract.tesseract_cmd):
            raise Exception(f"Tesseract not found at: {pytesseract.pytesseract.tesseract_cmd}")
        
        image = Image.open(file_path)
        # Convert image to RGB if it's not
        if image.mode != 'RGB':
            image = image.convert('RGB')
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        raise Exception(f"Image processing error: {str(e)}\nPlease verify Tesseract installation and path.")

def extract_from_txt(file_path):
    # Detect file encoding
    with open(file_path, 'rb') as file:
        raw_data = file.read()
        result = chardet.detect(raw_data)
        encoding = result['encoding']

    with open(file_path, 'r', encoding=encoding) as file:
        return file.read().strip()

def extract_text_from_file(file_path, file_type):
    try:
        if file_type == 'pdf':
            return extract_from_pdf(file_path)
        elif file_type in ['doc', 'docx']:
            return extract_from_docx(file_path)
        elif file_type == 'xlsx':
            return extract_from_excel(file_path)
        elif file_type == 'csv':
            return extract_from_csv(file_path)
        elif file_type in ['jpeg', 'jpg', 'png']:
            return extract_from_image(file_path)
        elif file_type == 'txt':
            return extract_from_txt(file_path)
        else:
            return "Unsupported file type"
    except Exception as e:
        return f"Error processing file: {str(e)}"

@app.route('/extract-text', methods=['POST'])
def extract_text():
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'success': False, 'error': 'File type not supported'}), 400
    
    try:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            file.save(temp_file.name)
            file_type = file.filename.rsplit('.', 1)[1].lower()
            extracted_text = extract_text_from_file(temp_file.name, file_type)
        
        # Clean up
        os.unlink(temp_file.name)
        
        if extracted_text:
            return jsonify({
                'success': True,
                'text': extracted_text,
                'fileType': file_type
            })
        else:
            return jsonify({
                'success': False,
                'error': 'Could not extract text from the file'
            }), 400
            
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

if __name__ == '__main__':
    app.run(debug=True)